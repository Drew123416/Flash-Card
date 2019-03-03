# 1 Introduce what this program is and how it works
import os
import smtplib
from io import BytesIO
import requests
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import time
def main():
    
    print("Welcome to the Flash Card Creator")
    time.sleep(.6)
    print("Please fill out the info....")
    name = input("Name of file: ")
    file = open(name + ".docx", "w")
    append = open(name + ".docx" , "rb")
    document = Document()
    # 2 Then ask the user how many flash cards they would like to create
    amount = int(input("How many flash cards: "))
    times = 0

    # makes sure it only makes as many flashcards as they instructed
    flashCardNum = 1

    # 3 Then have them fill each one out
    while times < amount:
        times += 1
        title = input(f"Title of {flashCardNum} card: " )

        header = document.add_heading(title + "\n")
        
        # Appending data to the .docx file
        sections = document.sections
        section = sections[0]
        section.page_width = Inches(5)
        body = input("line: ")
        paragraph = document.add_paragraph(body + "\n")
        paragraph.style = 'List Bullet'
        # Gives the bullet point functionality
        def line():
            body = input("line: ")
            paragraph = document.add_paragraph(body + "\n")
            paragraph.style = 'List Bullet'
        addLine = input("Add another line yes/no: ")
        while addLine.casefold() == "yes":
            line()
            addLine = input("Add another line yes/no: ")

        addImg = input("Add image yes/no: ")
        
        # asks if you want an image each flash card
        if addImg.casefold() == "yes":
            image = input("Image url: ")
            response = requests.get(image)  # no need to add stream=True
            binary_img = BytesIO(response.content) 
            document.add_picture(binary_img, width=Inches(2))
            document.save(name + ".docx")
        else:
            print("Okay") 
            time.sleep(.3)
            # saves the document
            document.save(name + ".docx")

           
        flashCardNum += 1
        # 4 Finally save all of that data to a text file so they can print it out
        
    append.close()
    
    another = input("Would you like to make another set yes/no: ")
    while another.casefold() == "yes":
        main()
    else:
        # Does all of the email stuff
        thereEmail = input("Your Email: ")

        subject = "Flash Cards"
        # grammar checks
        if amount == 1:
            msg = f"Thanks for creating {amount} flashcard with the Flash Card Creator \n Made by: Flash Card Creator"
        else: 
            msg = f"Thanks for creating {amount} flashcards with the Flash Card Creator \n Made by: Flash Card Creator"
            
                
        message = 'Subject: {}\n\n{}'.format(subject, msg)

        mail = smtplib.SMTP('smtp.gmail.com',587)

        mail.ehlo()

        mail.starttls()

        mail.login('drewgadams2016@gmail.com', 'Drew2016')

        mail.sendmail('drewgadams2016@gmail.com',thereEmail, message)

        mail.close()

        print("Thanks so much for using Flash Card Creator!")
        time.sleep(1)
        exit()  
    
# 5 Thank them for using your program 

if __name__ == "__main__":
    main()